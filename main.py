from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from docx import Document
from io import BytesIO

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/extract")
async def extract_docx(file: UploadFile = File(...), client_id: str = Form(...)):
    try:
        contents = await file.read()
        document = Document(BytesIO(contents))

        extracted_text = []

        for para in document.paragraphs:
            if para.text.strip():
                extracted_text.append(para.text.strip())

        for table in document.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    extracted_text.append(" | ".join(row_text))

        full_text = "\n".join(extracted_text)

        return {
            "client_id": client_id,
            "text": full_text
        }

    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

            "text": text
        }

    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})
